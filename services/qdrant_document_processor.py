"""
Qdrant Document Processor Service - Story 27.2
Handles document processing pipeline for vector database
Supports PDF, HTML, DOC, DOCX with Docling extraction
Uses LangChain for chunking and OpenAI for embeddings
"""

import os
import sys
import logging
import hashlib
import json
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
from pathlib import Path
import tempfile
import uuid

# Configure logging
logger = logging.getLogger(__name__)

# Import OpenAI for embeddings
try:
    from openai import OpenAI
    HAS_OPENAI = True
except ImportError:
    logger.warning("OpenAI not installed. Install with: pip install openai")
    HAS_OPENAI = False

# Import Qdrant client
try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import PointStruct
    HAS_QDRANT = True
except ImportError:
    logger.warning("qdrant-client not installed. Install with: pip install qdrant-client")
    HAS_QDRANT = False

# Import Docling for document extraction
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import ConversionStatus
    HAS_DOCLING = True
    logger.info("Docling available for enhanced document extraction")
except ImportError:
    logger.warning("Docling not installed. Install with: pip install docling")
    HAS_DOCLING = False

# Import LangChain for intelligent chunking
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    HAS_LANGCHAIN = True
    logger.info("LangChain available for intelligent text chunking")
except ImportError:
    logger.warning("LangChain not installed. Install with: pip install langchain")
    HAS_LANGCHAIN = False

# Fallback text extraction imports
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Add parent directory to path for local imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from utils.qdrant_init import get_qdrant_client, COLLECTION_NAME


class QdrantDocumentProcessor:
    """
    Complete document processing pipeline for Qdrant vector database.
    Follows Story 27.2 requirements with Docling, LangChain, and OpenAI.
    """
    
    def __init__(self):
        """Initialize the document processor with required components."""
        # Initialize document converter
        self.doc_converter = None
        if HAS_DOCLING:
            try:
                self.doc_converter = DocumentConverter()
                logger.info("Docling converter initialized successfully")
            except Exception as e:
                logger.warning(f"Failed to initialize Docling: {e}")
        
        # Initialize text splitter
        self.text_splitter = None
        if HAS_LANGCHAIN:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=int(os.getenv("CHUNK_SIZE", "1000")),
                chunk_overlap=int(os.getenv("CHUNK_OVERLAP", "200")),
                separators=["\n\n", "\n", ". ", ", ", " ", ""],
                length_function=len
            )
            logger.info("LangChain text splitter initialized")
        
        # Initialize OpenAI client
        self.openai_client = None
        if HAS_OPENAI:
            api_key = os.getenv("OPENAI_API_KEY")
            if api_key:
                self.openai_client = OpenAI(api_key=api_key)
                logger.info("OpenAI client initialized")
            else:
                logger.warning("OPENAI_API_KEY not found in environment")
        
        # Initialize Qdrant client
        self.qdrant_client = get_qdrant_client() if HAS_QDRANT else None
        
        # Embedding model configuration
        self.embedding_model = os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-small")
    
    def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from document using Docling or fallback methods.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (text_content, metadata)
        """
        metadata = {
            "extraction_method": None,
            "page_count": 0,
            "tables_count": 0,
            "figures_count": 0,
            "extraction_timestamp": datetime.now().isoformat()
        }
        
        # Try Docling first for best extraction
        if self.doc_converter:
            try:
                logger.info(f"Extracting with Docling: {file_path}")
                result = self.doc_converter.convert(Path(file_path))
                
                if result.status == ConversionStatus.SUCCESS:
                    # Export to markdown for best text representation
                    text_content = result.document.export_to_markdown()
                    
                    # Extract metadata from Docling result
                    metadata["extraction_method"] = "docling"
                    if hasattr(result.document, 'pages'):
                        metadata["page_count"] = len(result.document.pages)
                    if hasattr(result.document, 'tables'):
                        metadata["tables_count"] = len(result.document.tables)
                    if hasattr(result.document, 'figures'):
                        metadata["figures_count"] = len(result.document.figures)
                    
                    logger.info(f"Docling extraction successful: {len(text_content)} chars")
                    return text_content, metadata
                else:
                    logger.warning(f"Docling conversion failed: {result.status}")
            except Exception as e:
                logger.warning(f"Docling extraction failed: {e}")
        
        # Fallback to basic extraction methods
        file_ext = Path(file_path).suffix.lower()
        text_content = ""
        
        if file_ext == '.pdf' and HAS_PDF:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata["page_count"] = len(pdf_reader.pages)
                    metadata["extraction_method"] = "pypdf2"
                    
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n"
                
                logger.info(f"PDF extraction successful: {len(text_content)} chars")
            except Exception as e:
                logger.error(f"PDF extraction failed: {e}")
                raise
        
        elif file_ext in ['.docx', '.doc'] and HAS_DOCX:
            try:
                doc = Document(file_path)
                metadata["extraction_method"] = "python-docx"
                metadata["page_count"] = len(doc.paragraphs) // 10  # Approximate
                
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                
                logger.info(f"DOCX extraction successful: {len(text_content)} chars")
            except Exception as e:
                logger.error(f"DOCX extraction failed: {e}")
                raise
        
        elif file_ext in ['.html', '.htm']:
            try:
                from bs4 import BeautifulSoup
                with open(file_path, 'r', encoding='utf-8') as file:
                    soup = BeautifulSoup(file.read(), 'html.parser')
                    text_content = soup.get_text()
                    metadata["extraction_method"] = "beautifulsoup"
                
                logger.info(f"HTML extraction successful: {len(text_content)} chars")
            except ImportError:
                # Fallback to simple HTML text extraction
                with open(file_path, 'r', encoding='utf-8') as file:
                    import re
                    html_content = file.read()
                    text_content = re.sub('<[^<]+?>', '', html_content)
                    metadata["extraction_method"] = "regex"
            except Exception as e:
                logger.error(f"HTML extraction failed: {e}")
                raise
        
        elif file_ext == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text_content = file.read()
                    metadata["extraction_method"] = "direct"
                
                logger.info(f"Text extraction successful: {len(text_content)} chars")
            except Exception as e:
                logger.error(f"Text extraction failed: {e}")
                raise
        
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        return text_content, metadata
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks using LangChain or fallback method.
        
        Args:
            text: Text content to chunk
            
        Returns:
            List of text chunks
        """
        if self.text_splitter:
            # Use LangChain's intelligent chunking
            chunks = self.text_splitter.split_text(text)
            logger.info(f"Created {len(chunks)} chunks with LangChain")
            return chunks
        else:
            # Fallback to simple chunking
            chunk_size = int(os.getenv("CHUNK_SIZE", "1000"))
            chunk_overlap = int(os.getenv("CHUNK_OVERLAP", "200"))
            
            chunks = []
            start = 0
            text_length = len(text)
            
            while start < text_length:
                end = start + chunk_size
                chunk = text[start:end]
                
                # Try to break at sentence boundary
                if end < text_length:
                    last_period = chunk.rfind('. ')
                    if last_period > chunk_size * 0.8:  # If period is in last 20%
                        end = start + last_period + 1
                        chunk = text[start:end]
                
                chunks.append(chunk)
                start = end - chunk_overlap if end < text_length else end
            
            logger.info(f"Created {len(chunks)} chunks with fallback method")
            return chunks
    
    def create_embedding(self, text: str) -> Optional[List[float]]:
        """
        Generate embedding vector using OpenAI API.
        
        Args:
            text: Text to embed
            
        Returns:
            Embedding vector or None if failed
        """
        if not self.openai_client:
            logger.warning("OpenAI client not available for embeddings")
            return None
        
        try:
            response = self.openai_client.embeddings.create(
                model=self.embedding_model,
                input=text
            )
            
            embedding = response.data[0].embedding
            logger.debug(f"Created embedding with {len(embedding)} dimensions")
            return embedding
            
        except Exception as e:
            logger.error(f"Failed to create embedding: {e}")
            return None
    
    def process_document(
        self, 
        file_path: str, 
        metadata: Dict[str, Any],
        progress_callback: Optional[callable] = None
    ) -> Dict[str, Any]:
        """
        Complete document processing pipeline.
        
        Args:
            file_path: Path to the document file
            metadata: Document metadata
            progress_callback: Optional callback for progress updates
            
        Returns:
            Processing result dictionary
        """
        result = {
            "status": "failed",
            "document_id": metadata.get("document_id", str(uuid.uuid4())),
            "chunks_processed": 0,
            "vectors_stored": 0,
            "extraction_metadata": {},
            "error": None,
            "processing_time": 0
        }
        
        start_time = datetime.now()
        
        try:
            # Step 1: Extract text
            if progress_callback:
                progress_callback("Extracting text from document...", 0.1)
            
            text_content, extraction_metadata = self.extract_text(file_path)
            result["extraction_metadata"] = extraction_metadata
            
            if not text_content or len(text_content.strip()) < 10:
                raise ValueError("No meaningful text extracted from document")
            
            # Step 2: Chunk text
            if progress_callback:
                progress_callback("Splitting text into chunks...", 0.3)
            
            chunks = self.chunk_text(text_content)
            result["chunks_processed"] = len(chunks)
            
            if not chunks:
                raise ValueError("No chunks created from document")
            
            # Step 3: Generate embeddings and prepare points
            if progress_callback:
                progress_callback("Generating embeddings...", 0.5)
            
            if not self.openai_client:
                raise ValueError("OpenAI client not available - cannot generate embeddings")
            
            if not self.qdrant_client:
                raise ValueError("Qdrant client not available - cannot store vectors")
            
            points = []
            for i, chunk in enumerate(chunks):
                # Update progress
                if progress_callback and i % 5 == 0:
                    progress = 0.5 + (0.4 * i / len(chunks))
                    progress_callback(f"Processing chunk {i+1}/{len(chunks)}...", progress)
                
                # Generate embedding
                embedding = self.create_embedding(chunk)
                if not embedding:
                    logger.warning(f"Failed to create embedding for chunk {i}")
                    continue
                
                # Create point for Qdrant (use UUID for point ID)
                point_id = str(uuid.uuid4())
                point = PointStruct(
                    id=point_id,
                    vector=embedding,
                    payload={
                        **metadata,
                        "chunk_text": chunk[:500],  # Store first 500 chars for preview
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "chunk_size": len(chunk),
                        "embedding_model": self.embedding_model,
                        "extraction_method": extraction_metadata.get("extraction_method"),
                        "processed_at": datetime.now().isoformat()
                    }
                )
                points.append(point)
            
            # Step 4: Store in Qdrant
            if progress_callback:
                progress_callback("Storing vectors in database...", 0.9)
            
            if points:
                # Batch upload to Qdrant
                batch_size = 100
                for i in range(0, len(points), batch_size):
                    batch = points[i:i+batch_size]
                    self.qdrant_client.upsert(
                        collection_name=COLLECTION_NAME,
                        points=batch
                    )
                
                result["vectors_stored"] = len(points)
                result["status"] = "success"
                logger.info(f"Successfully stored {len(points)} vectors in Qdrant")
            else:
                raise ValueError("No vectors generated from document")
            
            if progress_callback:
                progress_callback("Processing complete!", 1.0)
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            result["error"] = str(e)
            result["status"] = "failed"
            
            if progress_callback:
                progress_callback(f"Error: {str(e)}", 0)
        
        finally:
            # Calculate processing time
            result["processing_time"] = (datetime.now() - start_time).total_seconds()
        
        return result
    
    def delete_document(self, document_id: str) -> bool:
        """
        Delete all vectors for a document from Qdrant.
        
        Args:
            document_id: Document ID to delete
            
        Returns:
            True if successful
        """
        if not self.qdrant_client:
            logger.warning("Qdrant client not available")
            return False
        
        try:
            # Delete points by filter
            from qdrant_client.models import Filter, FieldCondition, MatchValue
            
            self.qdrant_client.delete(
                collection_name=COLLECTION_NAME,
                points_selector=Filter(
                    must=[
                        FieldCondition(
                            key="document_id",
                            match=MatchValue(value=document_id)
                        )
                    ]
                )
            )
            
            logger.info(f"Deleted vectors for document: {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete document {document_id}: {e}")
            return False
    
    def search_similar(
        self, 
        query_text: str, 
        limit: int = 5,
        filter_conditions: Optional[Dict] = None
    ) -> List[Dict[str, Any]]:
        """
        Search for similar documents using semantic search.
        
        Args:
            query_text: Text to search for
            limit: Maximum number of results
            filter_conditions: Optional filter conditions
            
        Returns:
            List of search results
        """
        if not self.openai_client or not self.qdrant_client:
            logger.warning("OpenAI or Qdrant client not available for search")
            return []
        
        try:
            # Generate query embedding
            query_embedding = self.create_embedding(query_text)
            if not query_embedding:
                return []
            
            # Build filter if provided
            qdrant_filter = None
            if filter_conditions:
                from qdrant_client.models import Filter, FieldCondition, MatchValue
                conditions = []
                
                for key, value in filter_conditions.items():
                    conditions.append(
                        FieldCondition(
                            key=key,
                            match=MatchValue(value=value)
                        )
                    )
                
                if conditions:
                    qdrant_filter = Filter(must=conditions)
            
            # Search in Qdrant
            results = self.qdrant_client.search(
                collection_name=COLLECTION_NAME,
                query_vector=query_embedding,
                limit=limit,
                query_filter=qdrant_filter
            )
            
            # Format results
            formatted_results = []
            for result in results:
                formatted_results.append({
                    "score": result.score,
                    "document_id": result.payload.get("document_id"),
                    "chunk_text": result.payload.get("chunk_text"),
                    "chunk_index": result.payload.get("chunk_index"),
                    "original_filename": result.payload.get("original_filename"),
                    "organization": result.payload.get("organization"),
                    "document_type": result.payload.get("document_type")
                })
            
            logger.info(f"Found {len(formatted_results)} similar documents")
            return formatted_results
            
        except Exception as e:
            logger.error(f"Search failed: {e}")
            return []


def get_processor_status() -> Dict[str, bool]:
    """
    Check the status of all required components.
    
    Returns:
        Dictionary with component availability status
    """
    return {
        "docling": HAS_DOCLING,
        "langchain": HAS_LANGCHAIN,
        "openai": HAS_OPENAI and bool(os.getenv("OPENAI_API_KEY")),
        "qdrant": HAS_QDRANT and bool(get_qdrant_client()),
        "pdf_support": HAS_PDF,
        "docx_support": HAS_DOCX
    }


# Module test
if __name__ == "__main__":
    # Check component status
    print("=" * 50)
    print("Qdrant Document Processor Status Check")
    print("=" * 50)
    
    status = get_processor_status()
    for component, available in status.items():
        status_icon = "✅" if available else "❌"
        print(f"{status_icon} {component}: {'Available' if available else 'Not Available'}")
    
    print("=" * 50)
    
    # Test processing if all components available
    if all([status["openai"], status["qdrant"]]):
        print("\nTesting document processing pipeline...")
        processor = QdrantDocumentProcessor()
        
        # Create a test text file
        test_file = "/tmp/test_document.txt"
        with open(test_file, "w") as f:
            f.write("This is a test document for the Qdrant vector database. " * 20)
        
        # Process the test document
        result = processor.process_document(
            test_file,
            {
                "document_id": f"test_{uuid.uuid4()}",
                "document_type": "test",
                "organization": "test_org",
                "original_filename": "test_document.txt"
            }
        )
        
        print(f"\nProcessing result: {result['status']}")
        print(f"Chunks processed: {result['chunks_processed']}")
        print(f"Vectors stored: {result['vectors_stored']}")
        if result['error']:
            print(f"Error: {result['error']}")
    else:
        print("\nCannot test processing - missing required components")
        print("Ensure OPENAI_API_KEY is set and Qdrant is configured")