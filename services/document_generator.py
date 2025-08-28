"""Document generator service for creating DOCX exports with proper Slovenian character support."""

import os
import re
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import logging

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Service for generating documents from procurement data."""
    
    def __init__(self):
        """Initialize the document generator."""
        self.temp_dir = tempfile.gettempdir()
    
    def generate_document(self, data: Dict[str, Any], format_type: str = "DOCX", 
                         include_summary: bool = False, 
                         summary_content: Optional[str] = None) -> str:
        """
        Generate a document in the specified format.
        
        Args:
            data: Procurement data dictionary
            format_type: Format type (only "DOCX" supported)
            include_summary: Whether to include AI summary
            summary_content: Optional AI-generated summary content
            
        Returns:
            Path to the generated document
        """
        if format_type != "DOCX":
            format_type = "DOCX"  # Force DOCX as it's the only supported format
        
        return self._generate_docx(data, include_summary, summary_content)
    
    def _generate_docx(self, data: Dict[str, Any], 
                      include_summary: bool = False,
                      summary_content: Optional[str] = None) -> str:
        """Generate a DOCX document with proper formatting."""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Javno naročilo', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add procurement ID and name
        if data.get('id'):
            doc.add_heading(f"Številka: {data['id']}", level=2)
        if data.get('naziv'):
            doc.add_heading(f"Naziv: {data['naziv']}", level=2)
        
        doc.add_paragraph()
        
        # Add AI summary if provided
        if include_summary and summary_content:
            doc.add_heading('Povzetek z AI priporočili', level=1)
            self._add_formatted_content(doc, summary_content)
            doc.add_page_break()
        
        # Add sections
        self._add_basic_info_section(doc, data)
        self._add_contracting_authority_section(doc, data)
        self._add_procurement_details_section(doc, data)
        self._add_criteria_section(doc, data)
        self._add_lots_section(doc, data)
        self._add_additional_info_section(doc, data)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"javno_narocilo_{data.get('id', 'novo')}_{timestamp}.docx"
        filepath = os.path.join(self.temp_dir, filename)
        
        doc.save(filepath)
        logger.info(f"Generated DOCX document: {filepath}")
        
        return filepath
    
    def _add_formatted_content(self, doc, content: str):
        """Add content with markdown formatting support."""
        if not content:
            return
        
        lines = content.split('\n')
        for line in lines:
            if not line.strip():
                doc.add_paragraph()
                continue
            
            # Handle headers
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            # Handle bullet points
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text_to_paragraph(p, text)
            # Handle numbered lists - manual formatting for proper markdown support
            elif re.match(r'^\s*\d+\.\s+', line):
                match = re.match(r'^(\s*\d+\.\s+)(.+)', line)
                if match:
                    number_part = match.group(1)
                    text_part = match.group(2)
                    p = doc.add_paragraph()
                    p.add_run(number_part)
                    self._add_formatted_text_to_paragraph(p, text_part)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(-0.25)
            else:
                # Regular paragraph with formatting
                p = doc.add_paragraph()
                self._add_formatted_text_to_paragraph(p, line)
    
    def _add_formatted_text_to_paragraph(self, paragraph, text: str):
        """Add text with markdown formatting to a paragraph."""
        # Pattern to match bold text
        pattern = r'\*\*([^*]+)\*\*'
        last_end = 0
        
        for match in re.finditer(pattern, text):
            # Add text before the match
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            
            # Add bold text
            bold_run = paragraph.add_run(match.group(1))
            bold_run.bold = True
            
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])
    
    def _add_basic_info_section(self, doc, data: Dict[str, Any]):
        """Add basic information section."""
        doc.add_heading('Osnovni podatki', level=1)
        
        info = []
        if data.get('vrsta_postopka'):
            info.append(f"Vrsta postopka: {data['vrsta_postopka']}")
        if data.get('datum_objave'):
            info.append(f"Datum objave: {data['datum_objave']}")
        if data.get('rok_za_oddajo'):
            info.append(f"Rok za oddajo ponudb: {data['rok_za_oddajo']}")
        if data.get('ocenjena_vrednost'):
            info.append(f"Ocenjena vrednost: {data['ocenjena_vrednost']} EUR")
        
        for item in info:
            doc.add_paragraph(item)
        
        doc.add_paragraph()
    
    def _add_contracting_authority_section(self, doc, data: Dict[str, Any]):
        """Add contracting authority section."""
        doc.add_heading('Naročnik', level=1)
        
        info = []
        if data.get('narocnik_naziv'):
            info.append(f"Naziv: {data['narocnik_naziv']}")
        if data.get('narocnik_naslov'):
            info.append(f"Naslov: {data['narocnik_naslov']}")
        if data.get('narocnik_posta'):
            info.append(f"Pošta: {data['narocnik_posta']}")
        if data.get('narocnik_drzava'):
            info.append(f"Država: {data['narocnik_drzava']}")
        if data.get('kontaktna_oseba'):
            info.append(f"Kontaktna oseba: {data['kontaktna_oseba']}")
        if data.get('email'):
            info.append(f"E-pošta: {data['email']}")
        if data.get('telefon'):
            info.append(f"Telefon: {data['telefon']}")
        
        for item in info:
            doc.add_paragraph(item)
        
        doc.add_paragraph()
    
    def _add_procurement_details_section(self, doc, data: Dict[str, Any]):
        """Add procurement details section."""
        doc.add_heading('Podrobnosti naročila', level=1)
        
        if data.get('opis_narocila'):
            doc.add_heading('Opis naročila', level=2)
            doc.add_paragraph(data['opis_narocila'])
        
        if data.get('cpv_kode'):
            doc.add_heading('CPV kode', level=2)
            if isinstance(data['cpv_kode'], list):
                for code in data['cpv_kode']:
                    doc.add_paragraph(f"• {code}", style='List Bullet')
            else:
                doc.add_paragraph(data['cpv_kode'])
        
        if data.get('kraj_izvedbe'):
            doc.add_heading('Kraj izvedbe', level=2)
            doc.add_paragraph(data['kraj_izvedbe'])
        
        doc.add_paragraph()
    
    def _add_criteria_section(self, doc, data: Dict[str, Any]):
        """Add selection criteria section."""
        doc.add_heading('Merila za izbor', level=1)
        
        # Get criteria from data
        criteria = data.get('merila', [])
        if criteria:
            if isinstance(criteria, str):
                doc.add_paragraph(criteria)
            elif isinstance(criteria, list):
                for i, criterion in enumerate(criteria, 1):
                    if isinstance(criterion, dict):
                        p = doc.add_paragraph()
                        p.add_run(f"{i}. {criterion.get('naziv', 'Merilo')}")
                        if criterion.get('utez'):
                            p.add_run(f" (utež: {criterion['utez']}%)")
                        if criterion.get('opis'):
                            doc.add_paragraph(f"   {criterion['opis']}")
                    else:
                        doc.add_paragraph(f"{i}. {criterion}")
        else:
            doc.add_paragraph("Merila niso določena.")
        
        doc.add_paragraph()
    
    def _add_lots_section(self, doc, data: Dict[str, Any]):
        """Add lots section if applicable."""
        lots = data.get('sklopi', [])
        if lots and isinstance(lots, list) and len(lots) > 0:
            doc.add_heading('Sklopi', level=1)
            
            for i, lot in enumerate(lots, 1):
                if isinstance(lot, dict):
                    doc.add_heading(f"Sklop {i}: {lot.get('naziv', 'Brez naziva')}", level=2)
                    
                    if lot.get('opis'):
                        doc.add_paragraph(f"Opis: {lot['opis']}")
                    if lot.get('ocenjena_vrednost'):
                        doc.add_paragraph(f"Ocenjena vrednost: {lot['ocenjena_vrednost']} EUR")
                    if lot.get('cpv_kode'):
                        doc.add_paragraph(f"CPV kode: {lot['cpv_kode']}")
                    
                    doc.add_paragraph()
        
        doc.add_paragraph()
    
    def _add_additional_info_section(self, doc, data: Dict[str, Any]):
        """Add additional information section."""
        doc.add_heading('Dodatne informacije', level=1)
        
        info = []
        if data.get('dostop_do_dokumentacije'):
            info.append(f"Dostop do dokumentacije: {data['dostop_do_dokumentacije']}")
        if data.get('rok_za_vprašanja'):
            info.append(f"Rok za vprašanja: {data['rok_za_vprašanja']}")
        if data.get('datum_odpiranja'):
            info.append(f"Datum odpiranja ponudb: {data['datum_odpiranja']}")
        if data.get('veljavnost_ponudbe'):
            info.append(f"Veljavnost ponudbe: {data['veljavnost_ponudbe']} dni")
        
        for item in info:
            doc.add_paragraph(item)
        
        if data.get('dodatne_informacije'):
            doc.add_heading('Opombe', level=2)
            doc.add_paragraph(data['dodatne_informacije'])
        
        # Add footer with generation date
        doc.add_paragraph()
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer.add_run(f"Dokument generiran: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        footer.runs[0].font.size = Pt(10)
        footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)