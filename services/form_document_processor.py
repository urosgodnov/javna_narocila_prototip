"""
Form Document AI Processor - Story 28.4
Bridges form documents with existing AI processing pipeline
Handles text extraction, chunking, embeddings, and vector storage
"""

import os
import json
import sqlite3
import asyncio
import logging
from typing import Optional, Dict, List, Tuple
from datetime import datetime
from pathlib import Path
from io import BytesIO

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Text extraction libraries
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    logger.warning("PyPDF2 not installed. PDF extraction will be limited.")
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    logger.warning("python-docx not installed. DOCX extraction will be limited.")
    HAS_DOCX = False

try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    logger.warning("PIL/pytesseract not installed. OCR will not be available.")
    HAS_OCR = False

# Try importing Docling for enhanced extraction
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import ConversionStatus
    HAS_DOCLING = True
    logger.info("Docling available for enhanced document extraction")
except ImportError:
    logger.warning("Docling not installed. Using basic extraction methods.")
    HAS_DOCLING = False

# Try importing LangChain for intelligent chunking
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    HAS_LANGCHAIN = True
    logger.info("LangChain available for intelligent text chunking")
except ImportError:
    logger.warning("LangChain not installed. Using basic chunking methods.")
    HAS_LANGCHAIN = False


class TextExtractor:
    """Enhanced text extraction for various file formats with Docling support"""
    
    def __init__(self):
        """Initialize extractor with Docling if available"""
        self.use_docling = HAS_DOCLING and os.getenv('USE_DOCLING', 'true').lower() == 'true'
        self.converter = None
        
        if self.use_docling:
            try:
                self.converter = DocumentConverter()
                logger.info("Docling document converter initialized for enhanced extraction")
            except Exception as e:
                logger.warning(f"Docling initialization failed, using fallback: {e}")
                self.use_docling = False
    
    def _extract_with_docling(self, file_path: str) -> Optional[str]:
        """Extract text using Docling with enhanced structure preservation"""
        if not self.converter:
            return None
            
        try:
            # Convert document with Docling
            result = self.converter.convert(file_path)
            
            # Check conversion status
            if result.status != ConversionStatus.SUCCESS:
                logger.warning(f"Docling conversion failed with status: {result.status}")
                return None
            
            # Build structured text from the document
            text_parts = []
            
            # Get the main document content
            if hasattr(result, 'document') and result.document:
                doc = result.document
                
                # Extract text with structure preservation
                for element in doc.elements:
                    if hasattr(element, 'text') and element.text:
                        # Add section markers for better structure
                        if hasattr(element, 'type'):
                            if element.type == 'heading':
                                text_parts.append(f"\n## {element.text}\n")
                            elif element.type == 'table':
                                text_parts.append(f"\n[TABLE]\n{element.text}\n[/TABLE]\n")
                            elif element.type == 'list':
                                text_parts.append(f"\n{element.text}")
                            else:
                                text_parts.append(element.text)
                        else:
                            text_parts.append(element.text)
            
            # Fallback to raw text if structure not available
            if not text_parts and hasattr(result, 'text'):
                text_parts.append(result.text)
            
            if text_parts:
                extracted_text = '\n\n'.join(text_parts)
                logger.info(f"Docling successfully extracted {len(extracted_text)} characters")
                return extracted_text
            
            return None
            
        except Exception as e:
            logger.warning(f"Docling extraction failed: {e}")
            return None
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF files"""
        if not HAS_PDF:
            return "[PDF extraction not available - install PyPDF2]"
        
        try:
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX files"""
        if not HAS_DOCX:
            return "[DOCX extraction not available - install python-docx]"
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    @staticmethod
    def extract_from_image(file_path: str) -> str:
        """Extract text from images using OCR"""
        if not HAS_OCR:
            return "[OCR not available - install PIL and pytesseract]"
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            return ""
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """Main method to extract text based on file type - now with Docling priority"""
        file_type = file_type.lower()
        
        # Try Docling first if enabled (works with multiple formats)
        if self.use_docling and file_type in ['.pdf', '.docx', '.doc', '.pptx', '.xlsx']:
            logger.info(f"Attempting Docling extraction for {file_type} file")
            docling_result = self._extract_with_docling(file_path)
            if docling_result:
                return docling_result
            logger.info("Docling extraction failed, falling back to traditional methods")
        
        # Fallback to traditional extraction methods
        if file_type in ['.pdf']:
            return self.extract_from_pdf(file_path)
        elif file_type in ['.docx', '.doc']:
            return self.extract_from_docx(file_path)
        elif file_type in ['.jpg', '.jpeg', '.png']:
            return self.extract_from_image(file_path)
        elif file_type in ['.txt', '.md']:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Error reading text file: {e}")
                return ""
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return ""


class DocumentChunker:
    """Intelligent document chunking with LangChain RecursiveCharacterTextSplitter"""
    
    def __init__(self):
        """Initialize chunker with LangChain if available"""
        self.use_smart_chunking = (
            HAS_LANGCHAIN and 
            os.getenv('USE_SMART_CHUNKING', 'true').lower() == 'true'
        )
        
        if self.use_smart_chunking:
            # Get configuration from environment
            self.chunk_size = int(os.getenv('CHUNK_SIZE', '1500'))
            self.chunk_overlap = int(os.getenv('CHUNK_OVERLAP', '200'))
            
            # Initialize RecursiveCharacterTextSplitter
            self.splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""],
                keep_separator=True,
                is_separator_regex=False
            )
            logger.info(f"Smart chunking initialized: size={self.chunk_size}, overlap={self.chunk_overlap}")
        else:
            logger.info("Using basic chunking (50k character limit)")
            self.chunk_size = 50000
            self.chunk_overlap = 0
    
    def create_chunks(self, text: str, metadata: Optional[Dict] = None) -> List[Dict]:
        """Create intelligent chunks from text"""
        if not text:
            return []
        
        if self.use_smart_chunking:
            try:
                # Use LangChain for intelligent chunking
                chunk_texts = self.splitter.split_text(text)
                
                # Format chunks with metadata
                chunks = []
                for i, chunk_text in enumerate(chunk_texts):
                    chunk = {
                        'chunk_index': i,
                        'chunk_text': chunk_text,
                        'chunk_size': len(chunk_text),
                        'metadata': {
                            'chunk_method': 'langchain_recursive',
                            'chunk_overlap': self.chunk_overlap if i > 0 else 0,
                            'total_chunks': len(chunk_texts),
                            'chunk_config': {
                                'size': self.chunk_size,
                                'overlap': self.chunk_overlap
                            },
                            **(metadata or {})
                        }
                    }
                    chunks.append(chunk)
                
                logger.info(f"Created {len(chunks)} smart chunks using LangChain")
                return chunks
                
            except Exception as e:
                logger.error(f"Smart chunking failed, using fallback: {e}")
                return self._create_basic_chunks(text, metadata)
        else:
            return self._create_basic_chunks(text, metadata)
    
    def _create_basic_chunks(self, text: str, metadata: Optional[Dict] = None) -> List[Dict]:
        """Fallback basic chunking (original implementation)"""
        max_size = 50000
        
        if len(text) <= max_size:
            # Single chunk for small documents
            return [{
                'chunk_index': 0,
                'chunk_text': text,
                'chunk_size': len(text),
                'metadata': {
                    'chunk_method': 'basic',
                    'total_chunks': 1,
                    **(metadata or {})
                }
            }]
        
        # Split into multiple chunks if needed
        chunks = []
        for i in range(0, len(text), max_size):
            chunk_text = text[i:i + max_size]
            chunks.append({
                'chunk_index': len(chunks),
                'chunk_text': chunk_text,
                'chunk_size': len(chunk_text),
                'metadata': {
                    'chunk_method': 'basic',
                    'total_chunks': -1,  # Unknown until complete
                    **(metadata or {})
                }
            })
        
        # Update total chunks
        for chunk in chunks:
            chunk['metadata']['total_chunks'] = len(chunks)
        
        logger.info(f"Created {len(chunks)} basic chunks (50k limit)")
        return chunks
    
    def estimate_chunk_count(self, text_length: int) -> int:
        """Estimate number of chunks that will be created"""
        if self.use_smart_chunking:
            # Account for overlap
            effective_chunk_size = self.chunk_size - self.chunk_overlap
            return max(1, (text_length + effective_chunk_size - 1) // effective_chunk_size)
        else:
            return max(1, (text_length + 49999) // 50000)


class FormDocumentProcessor:
    """Bridge between form documents and existing AI pipeline"""
    
    def __init__(self, db_path: str = 'mainDB.db'):
        self.db_path = db_path
        self.text_extractor = TextExtractor()
        self.document_chunker = DocumentChunker()
        
        # Import existing AI processor components
        try:
            from ui.ai_processor import DocumentProcessor, EmbeddingGenerator, VectorStoreManager
            self.doc_processor = DocumentProcessor()
            self.embedding_gen = EmbeddingGenerator()
            self.vector_store = VectorStoreManager()
            self.ai_available = True
        except Exception as e:
            logger.warning(f"AI processor not available: {e}")
            self.ai_available = False
    
    def process_form_document(self, form_doc_id: int) -> Tuple[bool, str]:
        """
        Process a single form document through AI pipeline
        Returns: (success, message)
        """
        try:
            # 1. Get form document details
            form_doc = self._get_form_document(form_doc_id)
            if not form_doc:
                return False, f"Form document {form_doc_id} not found"
            
            # 2. Check if file should be processed
            if not self._should_process_file(form_doc):
                self._update_status(form_doc_id, 'skipped')
                return True, "File type not suitable for AI processing"
            
            # 3. Update status to processing
            self._update_status(form_doc_id, 'processing')
            
            # 4. Extract text from document
            text = self.text_extractor.extract_text(
                form_doc['file_path'],
                form_doc['file_type']
            )
            
            if not text or len(text.strip()) < 100:
                self._update_status(form_doc_id, 'skipped')
                return True, "Insufficient text content for processing"
            
            # 5. Create entry in ai_documents table
            ai_doc_id = self._create_ai_document_entry(
                filename=form_doc['original_name'],
                file_path=form_doc['file_path'],
                file_type=form_doc['file_type'],
                tip_dokumenta='form_upload',
                metadata={
                    'form_id': form_doc.get('form_id'),
                    'field_name': form_doc.get('field_name'),
                    'organization_id': form_doc.get('organization_id'),
                    'form_document_id': form_doc_id,
                    'text_length': len(text)
                }
            )
            
            # 6. Process through AI pipeline if available
            if self.ai_available:
                # Process through AI pipeline with extracted text
                success = self._process_with_ai(ai_doc_id, text, form_doc)
                
                if success:
                    # Link form document to AI document
                    self._link_to_ai_document(form_doc_id, ai_doc_id)
                    self._update_status(form_doc_id, 'completed')
                    
                    # Update processing metrics
                    self._update_processing_metrics(form_doc_id, ai_doc_id)
                    
                    return True, f"Document processed successfully (AI doc: {ai_doc_id})"
                else:
                    self._update_status(form_doc_id, 'failed', 'AI processing failed')
                    return False, "AI processing failed"
            else:
                # Just extract and store text without AI
                self._store_extracted_text(ai_doc_id, text)
                self._link_to_ai_document(form_doc_id, ai_doc_id)
                self._update_status(form_doc_id, 'completed')
                return True, "Text extracted and stored (AI not available)"
                
        except Exception as e:
            logger.error(f"Error processing document {form_doc_id}: {e}")
            self._update_status(form_doc_id, 'failed', str(e))
            return False, str(e)
    
    def _should_process_file(self, form_doc: Dict) -> bool:
        """Check if file type should be processed"""
        processable_types = ['.pdf', '.doc', '.docx', '.txt', '.md']
        # Optionally process images if OCR is available
        if HAS_OCR:
            processable_types.extend(['.jpg', '.jpeg', '.png'])
        
        file_type = form_doc.get('file_type', '').lower()
        return file_type in processable_types
    
    def _get_form_document(self, form_doc_id: int) -> Optional[Dict]:
        """Get form document details from database"""
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Get document with its associations
            cursor.execute("""
                SELECT d.*, 
                       a.form_id, a.form_type, a.field_name, a.organization_id
                FROM form_documents d
                LEFT JOIN form_document_associations a ON d.id = a.form_document_id
                WHERE d.id = ? AND d.is_active = 1
                LIMIT 1
            """, (form_doc_id,))
            
            result = cursor.fetchone()
            return dict(result) if result else None
    
    def _create_ai_document_entry(self, **kwargs) -> int:
        """Create entry in ai_documents table"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            metadata_json = json.dumps(kwargs.get('metadata', {}))
            
            cursor.execute("""
                INSERT INTO ai_documents (
                    filename, file_path, file_type, tip_dokumenta,
                    tags, description, processing_status, metadata_json
                ) VALUES (?, ?, ?, ?, ?, ?, 'pending', ?)
            """, (
                kwargs['filename'],
                kwargs['file_path'],
                kwargs.get('file_type', ''),
                kwargs.get('tip_dokumenta', 'form_upload'),
                kwargs.get('tags', ''),
                kwargs.get('description', ''),
                metadata_json
            ))
            
            conn.commit()
            return cursor.lastrowid
    
    def _process_with_ai(self, ai_doc_id: int, text: str, form_doc: Dict) -> bool:
        """Process document through existing AI pipeline with form metadata"""
        try:
            # Create intelligent chunks using DocumentChunker
            chunks = self.document_chunker.create_chunks(text, {
                'ai_document_id': ai_doc_id,
                'form_id': form_doc.get('form_id'),
                'processing_mode': 'ai_enabled'
            })
            
            if not chunks:
                logger.warning(f"No chunks created for document {ai_doc_id}")
                return False
            
            # Adapt chunk format for existing pipeline (if needed)
            # The doc_processor expects chunks with 'text' key
            formatted_chunks = []
            for chunk in chunks:
                formatted_chunk = {
                    'chunk_index': chunk['chunk_index'],
                    'text': chunk['chunk_text'],  # Map chunk_text to text
                    'chunk_size': chunk['chunk_size'],
                    'metadata': chunk.get('metadata', {})
                }
                formatted_chunks.append(formatted_chunk)
            
            # Generate embeddings
            texts = [chunk['text'] for chunk in formatted_chunks]
            embeddings = self.embedding_gen.generate_embeddings(texts)
            
            if not embeddings:
                logger.warning(f"No embeddings generated for document {ai_doc_id}")
                return False
            
            # Store in vector database with form metadata
            self._store_form_vectors(formatted_chunks, embeddings, ai_doc_id, form_doc)
            
            # Save chunks to database using our own method
            self._save_chunks_to_db(formatted_chunks, ai_doc_id)
            
            # Update AI document status
            self._update_ai_document_status(ai_doc_id, 'completed', len(chunks), len(embeddings))
            
            return True
            
        except Exception as e:
            logger.error(f"AI processing error for document {ai_doc_id}: {e}")
            self._update_ai_document_status(ai_doc_id, 'failed', error=str(e))
            return False
    
    def _store_form_vectors(self, chunks: List[Dict], embeddings: List[List[float]], 
                           ai_doc_id: int, form_doc: Dict):
        """Store vectors with form-specific metadata"""
        if not self.vector_store or not self.vector_store.client:
            logger.warning("Vector store not available")
            return
        
        import uuid
        from qdrant_client.models import PointStruct
        
        points = []
        for chunk, embedding in zip(chunks, embeddings):
            point_id = str(uuid.uuid4())
            
            # Enhanced payload with form metadata
            payload = {
                'document_id': ai_doc_id,
                'chunk_index': chunk['chunk_index'],
                'text': chunk['text'],
                'tip_dokumenta': 'form_upload',
                'form_id': form_doc.get('form_id'),
                'organization_id': form_doc.get('organization_id'),
                'field_name': form_doc.get('field_name'),
                'original_name': form_doc.get('original_name'),
                'form_type': form_doc.get('form_type', 'draft')
            }
            
            points.append(
                PointStruct(
                    id=point_id,
                    vector=embedding,
                    payload=payload
                )
            )
            
            # Store vector_id in chunk for reference
            chunk['vector_id'] = point_id
        
        # Batch upsert to Qdrant
        self.vector_store.client.upsert(
            collection_name=self.vector_store.collection_name,
            points=points
        )
        
        logger.info(f"Stored {len(points)} vectors for form document {ai_doc_id}")
    
    def _save_chunks_to_db(self, chunks: List[Dict], ai_doc_id: int):
        """Save chunks to database"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            for chunk in chunks:
                cursor.execute("""
                    INSERT INTO ai_document_chunks 
                    (document_id, chunk_index, chunk_text, chunk_size, vector_id)
                    VALUES (?, ?, ?, ?, ?)
                """, (
                    ai_doc_id,
                    chunk['chunk_index'],
                    chunk['text'],
                    chunk.get('chunk_size', len(chunk['text'])),
                    chunk.get('vector_id', '')
                ))
            
            conn.commit()
            logger.info(f"Saved {len(chunks)} chunks to database for document {ai_doc_id}")
    
    def _store_extracted_text(self, ai_doc_id: int, text: str):
        """Store extracted text with intelligent chunking"""
        # Create chunks using DocumentChunker
        chunks = self.document_chunker.create_chunks(text, {
            'ai_document_id': ai_doc_id,
            'processing_mode': 'text_only'
        })
        
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Store each chunk
            for chunk in chunks:
                cursor.execute("""
                    INSERT INTO ai_document_chunks 
                    (document_id, chunk_index, chunk_text, chunk_size, vector_id)
                    VALUES (?, ?, ?, ?, '')
                """, (
                    ai_doc_id, 
                    chunk['chunk_index'],
                    chunk['chunk_text'],
                    chunk['chunk_size']
                ))
            
            # Update chunk count in ai_documents
            if chunks:
                cursor.execute("""
                    UPDATE ai_documents 
                    SET chunk_count = ?
                    WHERE id = ?
                """, (len(chunks), ai_doc_id))
            
            conn.commit()
            
        logger.info(f"Stored {len(chunks)} chunks for document {ai_doc_id}")
    
    def _link_to_ai_document(self, form_doc_id: int, ai_doc_id: int):
        """Link form document to AI document"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE form_documents 
                SET ai_document_id = ?
                WHERE id = ?
            """, (ai_doc_id, form_doc_id))
            conn.commit()
    
    def _update_status(self, form_doc_id: int, status: str, error: str = None):
        """Update form document processing status"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE form_documents 
                SET processing_status = ?, 
                    processing_error = ?,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            """, (status, error, form_doc_id))
            conn.commit()
    
    def _update_ai_document_status(self, ai_doc_id: int, status: str, 
                                  chunk_count: int = 0, embedding_count: int = 0,
                                  error: str = None):
        """Update AI document status"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            if error:
                metadata = json.dumps({'error': error})
                cursor.execute("""
                    UPDATE ai_documents 
                    SET processing_status = ?, 
                        metadata_json = ?
                    WHERE id = ?
                """, (status, metadata, ai_doc_id))
            else:
                cursor.execute("""
                    UPDATE ai_documents 
                    SET processing_status = ?, 
                        chunk_count = ?
                    WHERE id = ?
                """, (status, chunk_count, ai_doc_id))
            
            conn.commit()
    
    def _update_processing_metrics(self, form_doc_id: int, ai_doc_id: int):
        """Update processing metrics for form document"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Get chunk and embedding counts
            cursor.execute("""
                SELECT COUNT(*) as chunk_count,
                       SUM(CASE WHEN vector_id != '' THEN 1 ELSE 0 END) as embedding_count
                FROM ai_document_chunks
                WHERE document_id = ?
            """, (ai_doc_id,))
            
            result = cursor.fetchone()
            if result:
                chunk_count, embedding_count = result
                
                # Store metrics in form document metadata
                cursor.execute("""
                    UPDATE form_documents
                    SET metadata_json = json_set(
                        COALESCE(metadata_json, '{}'),
                        '$.chunk_count', ?,
                        '$.embedding_count', ?
                    )
                    WHERE id = ?
                """, (chunk_count, embedding_count, form_doc_id))
                
                conn.commit()


class ProcessingQueue:
    """Async queue for document processing with retry logic"""
    
    def __init__(self, db_path: str = 'mainDB.db'):
        self.db_path = db_path
        self.processor = FormDocumentProcessor(db_path)
        self.max_retries = 3
        self.retry_delays = [5, 15, 30]  # seconds
    
    async def add_to_queue(self, form_doc_id: int, priority: int = 5):
        """Add document to processing queue"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT OR IGNORE INTO form_document_processing_queue
                (form_document_id, priority, status, retry_count, max_retries)
                VALUES (?, ?, 'pending', 0, ?)
            """, (form_doc_id, priority, self.max_retries))
            conn.commit()
        
        # Trigger async processing
        asyncio.create_task(self._process_with_retry(form_doc_id))
    
    async def _process_with_retry(self, form_doc_id: int):
        """Process document with retry logic"""
        for attempt in range(self.max_retries):
            try:
                # Update queue status
                self._update_queue_status(form_doc_id, 'processing', attempt)
                
                # Process document
                success, message = self.processor.process_form_document(form_doc_id)
                
                if success:
                    self._mark_complete(form_doc_id, message)
                    return
                else:
                    # If not the last attempt, retry
                    if attempt < self.max_retries - 1:
                        await asyncio.sleep(self.retry_delays[attempt])
                        continue
                    else:
                        self._mark_failed(form_doc_id, message)
                        
            except Exception as e:
                logger.error(f"Processing error for document {form_doc_id}: {e}")
                
                if attempt < self.max_retries - 1:
                    await asyncio.sleep(self.retry_delays[attempt])
                else:
                    self._mark_failed(form_doc_id, str(e))
    
    def _update_queue_status(self, form_doc_id: int, status: str, retry_count: int):
        """Update queue entry status"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE form_document_processing_queue
                SET status = ?, 
                    retry_count = ?,
                    started_at = CASE WHEN started_at IS NULL THEN CURRENT_TIMESTAMP ELSE started_at END
                WHERE form_document_id = ?
            """, (status, retry_count, form_doc_id))
            conn.commit()
    
    def _mark_complete(self, form_doc_id: int, message: str):
        """Mark queue entry as complete"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE form_document_processing_queue
                SET status = 'completed',
                    completed_at = CURRENT_TIMESTAMP,
                    error_message = ?
                WHERE form_document_id = ?
            """, (message, form_doc_id))
            conn.commit()
        
        logger.info(f"Document {form_doc_id} processing completed: {message}")
    
    def _mark_failed(self, form_doc_id: int, error: str):
        """Mark queue entry as failed"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE form_document_processing_queue
                SET status = 'failed',
                    completed_at = CURRENT_TIMESTAMP,
                    error_message = ?
                WHERE form_document_id = ?
            """, (error, form_doc_id))
            conn.commit()
        
        logger.error(f"Document {form_doc_id} processing failed: {error}")
    
    async def process_pending_queue(self):
        """Process all pending items in queue"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT form_document_id 
                FROM form_document_processing_queue
                WHERE status = 'pending'
                ORDER BY priority DESC, created_at ASC
            """)
            
            pending = cursor.fetchall()
        
        # Process all pending documents
        tasks = []
        for (form_doc_id,) in pending:
            tasks.append(self._process_with_retry(form_doc_id))
        
        if tasks:
            await asyncio.gather(*tasks)
            logger.info(f"Processed {len(tasks)} pending documents")


# Helper function for integration
def trigger_document_processing(form_doc_id: int):
    """Synchronous wrapper to trigger document processing"""
    try:
        # Create event loop if not exists
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
        
        # Create queue and add document
        queue = ProcessingQueue()
        loop.create_task(queue.add_to_queue(form_doc_id))
        
        logger.info(f"Document {form_doc_id} added to processing queue")
        return True
    except Exception as e:
        logger.error(f"Failed to trigger processing for document {form_doc_id}: {e}")
        return False